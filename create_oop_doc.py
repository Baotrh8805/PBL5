import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

doc = docx.Document()

# Adjust margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Colors
PRIMARY_COLOR = RGBColor(0, 86, 179) # #0056B3 (Deep Navy Blue)
SECONDARY_COLOR = RGBColor(0, 209, 178) # #00D1B2 (Teal)
TEXT_COLOR = RGBColor(30, 41, 59) # Slate Grey
CODE_COLOR = RGBColor(220, 38, 38) # Red-grey for code text

# Styles helper
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Inter'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    p.paragraph_format.space_after = Pt(20)

def add_heading_1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Inter'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(8)

def add_heading_2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Inter'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = SECONDARY_COLOR
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)

def add_paragraph(text, bold_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Inter'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = TEXT_COLOR
        
    run = p.add_run(text)
    run.font.name = 'Inter'
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_COLOR
    return p

def add_code_block(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(8)
    
    # Background grey shading
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
    pPr.append(shd)
    
    # Border
    pBdr = parse_xml(r'<w:pBdr {}><w:left w:val="single" w:sz="24" w:space="8" w:color="00D1B2"/></w:pBdr>'.format(nsdecls('w')))
    pPr.append(pBdr)
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(15, 23, 42)

# Content generation
add_title("BÁO CÁO PHÂN TÍCH TÍNH HƯỚNG ĐỐI TƯỢNG (OOP)\nTRONG ĐỀ TÀI ĐỒ ÁN PBL5")

add_paragraph("Đồ án tốt nghiệp / đồ án môn học PBL5 được phát triển trên nền tảng ngôn ngữ Java cùng Framework Spring Boot và Spring Data JPA. Java vốn là một ngôn ngữ thuần hướng đối tượng (Pure Object-Oriented Programming). Do đó, toàn bộ kiến trúc của hệ thống từ các lớp thực thể (Model/Entity), tầng truy cập dữ liệu (Repository), tầng xử lý logic (Service) đến tầng điều phối phản hồi (Controller) đều thỏa mãn và tuân thủ chặt chẽ 4 tính chất cốt lõi của lập trình hướng đối tượng (OOP): Tính đóng gói (Encapsulation), Tính kế thừa (Inheritance), Tính đa hình (Polymorphism) và Tính trừu tượng (Abstraction).")

add_heading_1("1. TÍNH ĐÓNG GÓI (ENCAPSULATION)")
add_paragraph("Tính đóng gói là kỹ thuật che giấu các thông tin chi tiết và trạng thái bên trong của một đối tượng, đồng thời chỉ cung cấp các phương thức công khai để truy cập và sửa đổi thông tin đó. Điều này giúp bảo vệ tính toàn vẹn của dữ liệu và tránh việc truy cập trái phép từ bên ngoài.")

add_heading_2("1.1. Sử dụng từ khóa truy cập private cho các thuộc tính")
add_paragraph("Trong các lớp thực thể như User, Post, Comment, toàn bộ các thuộc tính đại diện cho các trường thông tin trong cơ sở dữ liệu đều được khai báo là private. Ví dụ cụ thể trong lớp User.java:")

add_code_block('''public class User {
    @Id
    @GeneratedValue(strategy = GenerationType.IDENTITY)
    private Long id;

    @Column(unique = true, nullable = false)
    private String email;

    @Column(nullable = false, unique = true)
    private String fullName;

    private String password;
    
    @Enumerated(EnumType.STRING)
    private Role role;

    private Integer score = 0;
}''')

add_paragraph("Nhờ khai báo private, các lớp khác không thể tự ý thay đổi dữ liệu của đối tượng User một cách trực tiếp như user.score = -10 hoặc user.password = '123' (điều này có thể phá vỡ logic nghiệp vụ).")

add_heading_2("1.2. Cung cấp truy cập thông qua Getter và Setter công khai (Public)")
add_paragraph("Để thao tác với dữ liệu, các lớp bên ngoài phải gọi các phương thức Getter và Setter công khai. Điều này cho phép chúng ta thực hiện các bước kiểm tra logic hoặc định dạng lại dữ liệu trước khi lưu trữ hoặc phản hồi:")

add_code_block('''public String getEmail() {
    return email;
}

public void setEmail(String email) {
    // Chúng ta có thể bổ sung kiểm tra định dạng email hợp lệ tại đây trước khi gán
    this.email = email;
}

public Integer getScore() {
    return score;
}

public void setScore(Integer score) {
    if (score < 0) {
        throw new IllegalArgumentException("Điểm vi phạm không thể âm!");
    }
    this.score = score;
}''')

add_heading_2("1.3. Đóng gói tập giá trị cố định bằng Enum")
add_paragraph("Việc sử dụng các Enum như Role, UserStatus, PostStatus đóng gói các tập giá trị hợp lệ, ngăn chặn việc gán các chuỗi ký tự tự do làm sai lệch logic nghiệp vụ của hệ thống:")

add_code_block('''public enum Role {
    USER,
    MODERATOR,
    ADMIN
}''')

add_heading_1("2. TÍNH KẾ THỪA (INHERITANCE)")
add_paragraph("Kế thừa cho phép một lớp con sở hữu, mở rộng và tái sử dụng toàn bộ các thuộc tính và hành vi (phương thức) từ một lớp cha (superclass). Nó giúp giảm thiểu sự trùng lặp code, đồng nhất cấu trúc dữ liệu và xây dựng cấu trúc phân tầng rõ ràng giữa các lớp do lập trình viên định nghĩa.")

add_heading_2("2.1. Thiết kế Lớp Cha Trừu Tượng (Abstract Base Class) - BaseContent.java")
add_paragraph("Trong đồ án PBL5, các nội dung do người dùng tạo ra (User Generated Content) như Bài viết (Post) và Bình luận (Comment) đều chia sẻ chung những thuộc tính cơ bản như: Mã số định danh duy nhất (id) và Thời điểm khởi tạo nội dung (createdAt). Thay vì khai báo trùng lặp ở cả hai thực thể, đồ án đã trừu tượng hóa và đóng gói chúng vào một lớp cha chung có tên là BaseContent.java:")

add_code_block('''package com.pbl5.model;

import jakarta.persistence.*;
import java.time.LocalDateTime;

@MappedSuperclass
public abstract class BaseContent {

    @Id
    @GeneratedValue(strategy = GenerationType.IDENTITY)
    private Long id;

    @Column(nullable = false)
    private LocalDateTime createdAt;

    @PrePersist
    protected void onCreate() {
        this.createdAt = LocalDateTime.now();
    }

    public Long getId() {
        return id;
    }

    public void setId(Long id) {
        this.id = id;
    }

    public LocalDateTime getCreatedAt() {
        return createdAt;
    }

    public void setCreatedAt(LocalDateTime createdAt) {
        this.createdAt = createdAt;
    }
}''')

add_paragraph("Lớp cha BaseContent được khai báo abstract (lớp trừu tượng) vì hệ thống không bao giờ tạo trực tiếp đối tượng BaseContent độc lập, mà chỉ dùng nó làm khuôn mẫu thừa kế cho các thực thể cụ thể bên dưới.")

add_heading_2("2.2. Triển khai kế thừa từ Lớp Cha (Subclassing) - Post.java và Comment.java")
add_paragraph("Bằng cách sử dụng từ khóa extends trong Java, cả Post và Comment kế thừa trực tiếp từ BaseContent, tự động sở hữu toàn bộ các thuộc tính (id, createdAt) và các phương thức getter/setter liên quan:")

add_code_block('''// Thực thể Post kế thừa từ lớp cha BaseContent
@Entity
@Table(name = "posts")
public class Post extends BaseContent {
    
    @Column(columnDefinition = "TEXT")
    private String content;

    @Column(columnDefinition = "TEXT")
    private String imageUrl;

    // Các thuộc tính riêng của Post...
}''')

add_paragraph("Tương tự, lớp thực thể Comment.java cũng kế thừa từ BaseContent để chia sẻ chung mã định danh và ngày tạo:")

add_code_block('''// Thực thể Comment kế thừa từ lớp cha BaseContent
@Entity
@Table(name = "comments")
public class Comment extends BaseContent {

    @Column(nullable = false, columnDefinition = "TEXT")
    private String content;

    @ManyToOne(fetch = FetchType.LAZY)
    @JoinColumn(name = "post_id", nullable = false)
    private Post post;

    // Các thuộc tính riêng của Comment...
}''')

add_paragraph("Hiệu quả của việc kế thừa từ class BaseContent: Giúp loại bỏ hoàn toàn mã nguồn trùng lặp ở các lớp con (tránh việc khai báo lặp đi lặp lại các trường id, createdAt, các phương thức truy cập và bộ bắt sự kiện @PrePersist). Khi cần mở rộng các thông số ghi vết (ví dụ thêm trường thông tin updatedAt - ngày cập nhật), lập trình viên chỉ cần thêm duy nhất một lần tại lớp cha BaseContent, toàn bộ các lớp con Post và Comment sẽ ngay lập tức được kế thừa và áp dụng tự động.")

add_heading_1("3. TÍNH ĐA HÌNH (POLYMORPHISM)")
add_paragraph("Đa hình cho phép một đối tượng có thể thể hiện dưới nhiều dạng khác nhau. Trong lập trình, đa hình thể hiện qua hai dạng: Nạp chồng phương thức (Method Overloading) và Ghi đè phương thức (Method Overriding). Ngoài ra, đa hình còn thể hiện ở việc tham chiếu lớp con thông qua kiểu của lớp cha (Interface/Abstract class).")

add_heading_2("3.1. Ghi đè phương thức (Method Overriding)")
add_paragraph("Đa hình ghi đè được thể hiện rõ khi chúng ta triển khai giao diện hoặc kế thừa lớp cơ sở và ghi đè phương thức đó bằng chú thích @Override. Điển hình là việc tùy biến lớp xử lý xác thực người dùng:")

add_code_block('''@Override
public UserDetails loadUserByUsername(String email) throws UsernameNotFoundException {
    User user = userRepository.findByEmail(email)
            .orElseThrow(() -> new UsernameNotFoundException("Không tìm thấy email: " + email));
    return new org.springframework.security.core.userdetails.User(
            user.getEmail(),
            user.getPassword(),
            Collections.singletonList(new SimpleGrantedAuthority("ROLE_" + user.getRole().name()))
    );
}''')

add_paragraph("Phương thức loadUserByUsername ban đầu được khai báo trong interface UserDetailsService của Spring Security. Class CustomUserDetailsService của đồ án đã ghi đè phương thức này để tùy biến cách tải tài khoản người dùng từ PostgreSQL thông qua UserRepository.")

add_heading_2("3.2. Nạp chồng phương thức (Method Overloading)")
add_paragraph("Nạp chồng phương thức cho phép tạo nhiều phương thức cùng tên trong cùng một lớp nhưng khác biệt về số lượng hoặc kiểu tham số đầu vào. Ví dụ trong UserRepository.java:")

add_code_block('''// Tìm kiếm người dùng chỉ dựa trên email
Optional<User> findByEmail(String email);

// Tìm kiếm người dùng dựa trên cả email HOẶC username (nạp chồng tham số)
Optional<User> findByEmailOrUsername(String email, String username);''')

add_heading_2("3.3. Đa hình thông qua Interface (Dependency Injection)")
add_paragraph("Khi khai báo các dependency trong Controller, chúng ta sử dụng kiểu dữ liệu là Interface chứ không phải lớp cụ thể. Spring Boot IoC Container sẽ tự động tìm kiếm Bean triển khai tương ứng và tiêm (inject) vào lúc chạy:")

add_code_block('''@RestController
@RequestMapping("/api/moderator")
public class ModeratorController {
    
    // Tham chiếu đa hình thông qua interface UserRepository
    @Autowired
    private UserRepository userRepository; 
    
    // Tham chiếu đa hình thông qua interface PasswordEncoder
    @Autowired
    private PasswordEncoder passwordEncoder;
}''')

add_paragraph("Tại thời điểm chạy, hệ thống không quan tâm lớp triển khai cụ thể là gì, chỉ cần lớp đó thực thi đầy đủ các giao diện hợp đồng UserRepository và PasswordEncoder, điều này mang lại tính linh hoạt cực kỳ cao.")

add_heading_1("4. TÍNH TRỪU TƯỢNG (ABSTRACTION)")
add_paragraph("Tính trừu tượng giúp tập trung vào những hành vi cốt lõi của đối tượng (đối tượng làm cái gì) thay vì đi sâu vào chi tiết triển khai cụ thể (đối tượng làm thế nào). Nó được biểu diễn qua các Giao diện (Interface) và Lớp trừu tượng (Abstract class).")

add_heading_2("4.1. Trừu tượng hóa thao tác cơ sở dữ liệu qua Repository Interface")
add_paragraph("UserRepository, PostRepository và ReportRepository đều là các interface trừu tượng. Chúng ta chỉ cần khai báo phương thức và các tham số truyền vào:")

add_code_block('''public interface ReportRepository extends JpaRepository<Report, Long> {
    List<Report> findByStatus(ReportStatus status);
}''')

add_paragraph("Lập trình viên hoàn toàn không cần quan tâm đến cách viết câu lệnh SQL 'SELECT * FROM reports WHERE status = ?' hay cách kết nối driver JDBC thế nào. Tất cả chi tiết triển khai phức tạp này đã được trừu tượng hóa và ẩn đi bởi Spring Data JPA.")

add_heading_2("4.2. Trừu tượng hóa nghiệp vụ thông qua tầng Service và DTO")
add_paragraph("Việc sử dụng các lớp Data Transfer Object (DTO) như PostDTO, CommentDTO giúp trừu tượng hóa các thực thể dữ liệu phức tạp trong Database thành các cấu trúc dữ liệu gọn nhẹ, vừa vặn chỉ chứa các thông tin cần thiết phục vụ cho việc truyền tải qua giao thức HTTP API.")

add_heading_1("KẾT LUẬN")
add_paragraph("Đồ án PBL5 là minh chứng điển hình cho việc áp dụng lập trình hướng đối tượng (OOP) một cách mẫu mực và thực tế. Nhờ tuân thủ nghiêm ngặt 4 nguyên lý OOP, mã nguồn của đồ án đạt được các tiêu chuẩn:")
add_paragraph("- Dễ dàng bảo trì (Maintainable): Thay đổi cấu trúc dữ liệu bên trong không ảnh hưởng đến lớp gọi nhờ tính Đóng gói.")
add_paragraph("- Tái sử dụng cao (Reusable): Tiết kiệm thời gian lập trình nhờ tính Kế thừa từ lớp cha BaseContent do hệ thống tự thiết kế.")
add_paragraph("- Dễ dàng mở rộng (Extensible): Khả năng thay đổi linh hoạt logic kiểm duyệt hoặc thêm các nhà cung cấp đăng nhập (Google, Facebook) nhờ tính Đa hình và Trừu tượng.")

doc.save("PBL5_OOP_Analysis.docx")
print("Successfully generated Word document PBL5_OOP_Analysis.docx with custom class inheritance details.")
